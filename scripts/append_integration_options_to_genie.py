"""Append an 'Integration Path Options' section to Genie_Platform_Documentation.docx.

Preserves the original document and writes the output to
Genie_Platform_Documentation_Updated.docx. The new content uses the exact font
(Arial 11pt) and colour palette (#0F172A for H1, #1A56DB for H2, #1F2937 for H3,
#F9FAFB for zebra striping, #0F172A with white text for table headers) that the
original document already uses, so the appended pages are visually indistinguishable.

Run:  python scripts/append_integration_options_to_genie.py
"""
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# Colours sampled from the existing Genie doc
COLOR_H1        = RGBColor(0x0F, 0x17, 0x2A)
COLOR_H2        = RGBColor(0x1A, 0x56, 0xDB)
COLOR_H3        = RGBColor(0x1F, 0x29, 0x37)
COLOR_BODY      = RGBColor(0x1F, 0x29, 0x37)
COLOR_MUTED     = RGBColor(0x4B, 0x55, 0x63)
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
FILL_HEADER_HEX = "0F172A"
FILL_ZEBRA_HEX  = "F9FAFB"
FILL_ACCENT_HEX = "EFF6FF"
BORDER_HEX      = "E5E7EB"


def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color=BORDER_HEX, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def _style_run(run, *, bold=False, italic=False, size=11, color=COLOR_BODY, font_name="Arial"):
    run.font.name = font_name
    # Ensure East-Asian + complex-script fonts also map to Arial so the doc is visually uniform
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _style_cell(cell, text, *, bold=False, size=10, color=COLOR_BODY, shading=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    _style_run(run, bold=bold, size=size, color=color)
    if shading:
        _set_cell_shading(cell, shading)
    _set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    for i, h in enumerate(headers):
        _style_cell(
            table.rows[0].cells[i], h,
            bold=True, size=10, color=COLOR_WHITE, shading=FILL_HEADER_HEX,
        )

    for r, row in enumerate(rows, start=1):
        zebra = FILL_ZEBRA_HEX if r % 2 == 0 else None
        for c, val in enumerate(row):
            _style_cell(table.rows[r].cells[c], val, bold=False, size=10, shading=zebra)

    # blank line after table
    doc.add_paragraph()
    return table


def _add_heading(doc, text, *, level=1):
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text)
        _style_run(run, bold=True, size=16, color=COLOR_H1)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run = p.add_run(text)
        _style_run(run, bold=True, size=13, color=COLOR_H2)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    else:
        run = p.add_run(text)
        _style_run(run, bold=True, size=11, color=COLOR_H3)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def _add_para(doc, text, *, bold=False, italic=False, size=11, color=COLOR_BODY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _style_run(run, bold=bold, italic=italic, size=size, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p


_NUMBERED_COUNTERS: dict[int, int] = {}


def _add_bullet(doc, text, *, size=11):
    p = doc.add_paragraph()
    run = p.add_run("• " + text)
    _style_run(run, size=size, color=COLOR_BODY)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    return p


def _reset_numbered(group: int = 0):
    _NUMBERED_COUNTERS[group] = 0


def _add_numbered(doc, text, *, size=11, group: int = 0):
    _NUMBERED_COUNTERS[group] = _NUMBERED_COUNTERS.get(group, 0) + 1
    n = _NUMBERED_COUNTERS[group]
    p = doc.add_paragraph()
    run = p.add_run(f"{n}.  {text}")
    _style_run(run, size=size, color=COLOR_BODY)
    p.paragraph_format.left_indent = Cm(0.85)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_kv(doc, key, value):
    """A crisp 'Label — value' line: bold key, regular value, one line."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(2)
    k = p.add_run(f"{key}: ")
    _style_run(k, bold=True, size=11, color=COLOR_H3)
    v = p.add_run(value)
    _style_run(v, size=11, color=COLOR_BODY)
    return p


def append_section(doc):
    _add_page_break(doc)

    # ===== Section title =====
    _add_heading(doc, "Section 5: Integration Path Options", level=1)
    _add_para(
        doc,
        "Two paths exist to take Genie from PoC to production. This section lists the facts — "
        "models, tech stack, effort, limitations, prerequisites — for each, and states the "
        "recommended approach.",
    )

    # ===== 5.1 Snapshot =====
    _add_heading(doc, "5.1  Snapshot", level=2)
    _add_table(
        doc,
        headers=["Attribute", "Option A — GitHub Copilot API", "Option B — Direct Model APIs"],
        rows=[
            ["Access method", "Undocumented endpoint (api.githubcopilot.com)", "Official provider APIs (Anthropic, OpenAI, Azure OpenAI)"],
            ["Licensing", "No public third-party API (verified from GitHub docs)", "Standard commercial contracts"],
            ["Cost model", "Per-user Copilot seat", "Per-token usage"],
            ["Model freedom", "Restricted to Copilot's catalogue", "Any licensed model"],
            ["SLA", "None for third-party use", "Enterprise SLAs available"],
            ["Effort to production", "4–6 months (legal-blocked)", "4–6 weeks"],
            ["Fit for enterprise sale", "Not suitable", "Suitable"],
        ],
        col_widths=[3.8, 5.6, 6.1],
    )

    # ===== 5.2 Option A =====
    _add_heading(doc, "5.2  Option A — GitHub Copilot API", level=2)

    _add_heading(doc, "5.2.1  Key Facts", level=3)
    _add_kv(doc, "Access method", "OAuth device-code flow → session token → POST to api.githubcopilot.com/chat/completions")
    _add_kv(doc, "Licensing", "No public API. GitHub officially supports only MCP servers and Copilot Extensions for third parties.")
    _add_kv(doc, "Models available", "GPT-4o, GPT-4.1, GPT-5.x, Claude Sonnet 3.7/4.x, Gemini 2.0, o3-mini (Copilot catalogue only)")
    _add_kv(doc, "Authentication", "GitHub OAuth; every end-user needs a Copilot Business/Enterprise seat")
    _add_kv(doc, "Cost model", "Copilot seat per developer (paid to GitHub)")
    _add_kv(doc, "SLA / support", "None for third-party use of the endpoint")
    _add_kv(doc, "Data residency", "Controlled by GitHub, not configurable by client")

    _add_heading(doc, "5.2.2  Tech Stack", level=3)
    _add_bullet(doc, "VS Code extension — TypeScript 5.5, @vscode/vsce")
    _add_bullet(doc, "Agent runtime — Python 3.11+, LangGraph 0.2+, LangChain 0.3+")
    _add_bullet(doc, "LLM client — langchain-openai against Copilot base URL with editor-identification headers")
    _add_bullet(doc, "Auth — GitHub OAuth device-code flow")
    _add_bullet(doc, "Transport — HTTPS with corporate proxy and corporate CA bundle")

    _add_heading(doc, "5.2.3  Limitations", level=3)
    _add_bullet(doc, "No public API contract — GitHub can revoke access, change headers, or throttle at will.")
    _add_bullet(doc, "Terms of Service risk — reselling Copilot output needs a partnership agreement.")
    _add_bullet(doc, "Undocumented rate limits — unpredictable throttling under agentic load.")
    _add_bullet(doc, "Locked to Copilot's model catalogue — no Claude Opus 4.7, no o3/o4-mini, no fine-tuning.")
    _add_bullet(doc, "No usage telemetry exposed — client cannot produce audit trail or cost reports.")
    _add_bullet(doc, "Data residency fixed to GitHub — blocks clients with EU/UK/APAC-only mandates.")
    _add_bullet(doc, "Short-lived session tokens — continuous refresh required.")

    _add_heading(doc, "5.2.4  Prerequisites", level=3)
    _add_bullet(doc, "Legal — partnership / licensing agreement with GitHub or Microsoft.")
    _add_bullet(doc, "Procurement — Copilot Business/Enterprise licence for every end-user.")
    _add_bullet(doc, "Compliance — DPIA covering prompt content and GitHub data routing.")
    _add_bullet(doc, "Network — corporate proxy allow-list and CA bundle.")
    _add_bullet(doc, "Monitoring — client-side usage, latency, and error dashboards (no telemetry from GitHub).")

    _add_heading(doc, "5.2.5  Effort to Production", level=3)
    _add_table(
        doc,
        headers=["Workstream", "Effort"],
        rows=[
            ["Partnership negotiation with GitHub/Microsoft", "3–6 months (legal, blocking)"],
            ["Engineering hardening (tokens, retries, observability)", "4 weeks"],
            ["Compliance documentation (DPIA, audit)", "2 weeks"],
            ["Total", "4–6 months"],
        ],
        col_widths=[10.0, 5.5],
    )

    _add_heading(doc, "5.2.6  Pros / Cons", level=3)
    _add_para(doc, "Pros", bold=True)
    _add_bullet(doc, "PoC already works — zero migration.")
    _add_bullet(doc, "Single developer identity (GitHub OAuth).")
    _add_bullet(doc, "No separate LLM billing — covered by Copilot seat.")
    _add_para(doc, "Cons", bold=True)
    _add_bullet(doc, "Not commercially licensable for enterprise sale.")
    _add_bullet(doc, "Cannot select best model per task.")
    _add_bullet(doc, "No SLA, no support, no audit trail.")
    _add_bullet(doc, "Blocks regulated / data-residency clients.")

    # ===== 5.3 Option B =====
    _add_heading(doc, "5.3  Option B — Direct Model APIs (Claude + OpenAI)", level=2)

    _add_heading(doc, "5.3.1  Key Facts", level=3)
    _add_kv(doc, "Access method", "Direct calls to Anthropic, OpenAI, and Azure OpenAI endpoints via a centralised LLM gateway")
    _add_kv(doc, "Licensing", "Standard commercial API contracts with each provider")
    _add_kv(doc, "Models available", "Full catalogue: Claude Opus 4.7, Sonnet 4.6, Haiku 4.5; GPT-5.5, GPT-5 mini, o3, o4-mini")
    _add_kv(doc, "Authentication", "Provider API keys held centrally in Key Vault; gateway issues scoped tokens to users")
    _add_kv(doc, "Cost model", "Pay-per-token; per-user and per-day caps enforceable at gateway")
    _add_kv(doc, "SLA / support", "Enterprise tier with Anthropic, OpenAI, and Azure OpenAI")
    _add_kv(doc, "Data residency", "US / EU / UK / APAC selectable via Azure OpenAI")

    _add_heading(doc, "5.3.2  Model → Action Mapping", level=3)
    _add_table(
        doc,
        headers=["Action / Review Type", "Recommended Model", "Reason"],
        rows=[
            [
                "Quality, Performance, OrgStd, CK Design review",
                "Claude Sonnet 4.6",
                "Best tool-calling reliability; 0.5–1.5s first-token latency; one-fifth flagship cost.",
            ],
            [
                "Security review",
                "Claude Opus 4.7",
                "Lowest hallucination rate; strongest on subtle security reasoning.",
            ],
            [
                "Guided Apply (whole-file edit)",
                "Claude Opus 4.7",
                "Whole-file edits are where cheaper models hallucinate — Opus is the safe choice.",
            ],
            [
                "Deep Analysis mode (user-triggered)",
                "OpenAI o3",
                "Top SWE-bench Verified score; deep reasoning for hard bugs; acceptable for opt-in actions.",
            ],
            [
                "Fast syntax review, relevance gate",
                "Claude Haiku 4.5",
                "Fast, inexpensive, sufficient for yes/no classification and simple formatting.",
            ],
            [
                "Assistant chat (Explain, Refactor, Docstring, Unit Test)",
                "Claude Sonnet 4.6 or GPT-5.5",
                "Responsive streaming; strong code quality; good cost–quality balance for interactive use.",
            ],
            [
                "Fallback / degraded mode",
                "GPT-5 mini",
                "Cheap, fast, widely available — used when primary provider is down or over budget.",
            ],
        ],
        col_widths=[4.8, 4.0, 6.7],
    )

    _add_heading(doc, "5.3.3  Tech Stack", level=3)
    _add_bullet(doc, "VS Code extension — TypeScript 5.5, @vscode/vsce (unchanged)")
    _add_bullet(doc, "Agent runtime — Python 3.11+, LangGraph 0.2+, LangChain 0.3+ (unchanged)")
    _add_bullet(doc, "LLM clients — langchain-anthropic, langchain-openai")
    _add_bullet(doc, "LLM gateway — LiteLLM or Portkey (unified OpenAI-compatible endpoint, routing, quotas, cost tracking)")
    _add_bullet(doc, "Secrets — Azure Key Vault / AWS Secrets Manager / HashiCorp Vault")
    _add_bullet(doc, "Observability — OpenTelemetry + Grafana (latency, tokens, cost)")
    _add_bullet(doc, "Transport — HTTPS with corporate proxy and CA bundle (unchanged)")

    _add_heading(doc, "5.3.4  Limitations", level=3)
    _add_bullet(doc, "Pay-per-token cost — scales with usage; capped via gateway.")
    _add_bullet(doc, "Requires key-distribution infrastructure (gateway + Key Vault).")
    _add_bullet(doc, "Provider outages — mitigated by dual-provider fallback.")
    _add_bullet(doc, "Multiple vendor contracts and DPAs instead of a single Copilot seat.")

    _add_heading(doc, "5.3.5  Prerequisites", level=3)
    _add_bullet(doc, "Commercial — paid accounts with Anthropic and OpenAI; Azure OpenAI if region pinning is required.")
    _add_bullet(doc, "Infrastructure — LLM gateway (LiteLLM / Portkey) and secrets management.")
    _add_bullet(doc, "Compliance — signed DPAs with each provider; zero-retention enrolment.")
    _add_bullet(doc, "Observability — dashboards for latency, tokens, cost per user / per task.")
    _add_bullet(doc, "Engineering — swap the LLM client layer (LangChain already provides the abstraction).")

    _add_heading(doc, "5.3.6  Effort to Production", level=3)
    _add_table(
        doc,
        headers=["Workstream", "Effort"],
        rows=[
            ["Replace Copilot client with provider-agnostic LangChain factory", "3–5 days"],
            ["LLM gateway (LiteLLM) + Key Vault wiring", "1 week"],
            ["Observability — dashboards, cost tracking, alerts", "1 week"],
            ["Cross-provider fallback and error handling", "3 days"],
            ["Per-task model routing configuration", "3 days"],
            ["DPAs + zero-retention enrolment (parallel with engineering)", "2–4 weeks"],
            ["Total", "4–6 weeks"],
        ],
        col_widths=[10.0, 5.5],
    )

    _add_heading(doc, "5.3.7  Pros / Cons", level=3)
    _add_para(doc, "Pros", bold=True)
    _add_bullet(doc, "Fully licensed — commercially defensible for enterprise sale.")
    _add_bullet(doc, "Best model per task — measurable lift in review accuracy.")
    _add_bullet(doc, "Enterprise SLAs, audit trail, data residency under client control.")
    _add_bullet(doc, "Gateway enforces cost caps and quotas.")
    _add_bullet(doc, "New models pluggable without re-architecture.")
    _add_para(doc, "Cons", bold=True)
    _add_bullet(doc, "One-time migration of the LLM layer (~1 week engineering).")
    _add_bullet(doc, "Cost scales linearly with usage (mitigated by caps).")
    _add_bullet(doc, "Multiple vendor relationships to manage.")

    # ===== 5.4 Decision matrix =====
    _add_heading(doc, "5.4  Decision Matrix", level=2)
    _add_table(
        doc,
        headers=["Factor", "Option A", "Option B", "Winner"],
        rows=[
            ["Legal certainty", "No public API", "Commercial contracts", "B"],
            ["Best model per task", "No", "Yes", "B"],
            ["SLA and support", "None", "Enterprise tier", "B"],
            ["Audit trail / data residency", "No", "Yes", "B"],
            ["Engineering effort", "Minimal", "3–4 weeks", "A"],
            ["Time to enterprise deal", "4–6 months", "4–6 weeks", "B"],
            ["Recurring cost, light use", "Full seat", "Lower (per-token)", "B"],
            ["Recurring cost, heavy use", "Capped at seat", "Scales linearly", "A"],
            ["Future model flexibility", "None", "Full", "B"],
        ],
        col_widths=[6.5, 3.7, 3.8, 1.5],
    )

    # ===== 5.5 Recommended approach =====
    _add_heading(doc, "5.5  Recommended Approach", level=2)
    _add_para(
        doc,
        "Proceed with Option B — Direct Model APIs behind a centralised LLM gateway. Default review "
        "routing to Claude Sonnet 4.6; route Security and Guided Apply to Claude Opus 4.7; offer "
        "OpenAI o3 as a user-triggered Deep Analysis action. The existing LangGraph agent, tool "
        "integrations, and VS Code surface remain unchanged.",
        bold=True,
    )

    _add_heading(doc, "5.5.1  Delivery Plan", level=3)
    _add_table(
        doc,
        headers=["Phase", "Duration", "Deliverable"],
        rows=[
            ["1 — LLM abstraction", "Week 1–2", "Replace Copilot client with LangChain provider factory; add Anthropic and OpenAI drivers."],
            ["2 — Gateway + secrets", "Week 3", "Deploy LiteLLM gateway; move provider keys to Key Vault; gateway enforces per-user quotas."],
            ["3 — Observability + routing", "Week 4", "Grafana dashboards; per-task model routing; cross-provider fallback."],
            ["Parallel — Legal", "Week 1–4", "DPAs with Anthropic and OpenAI; zero-retention enrolment; region selection."],
        ],
        col_widths=[4.0, 2.5, 9.0],
    )

    _add_heading(doc, "5.5.2  Production Default Routing", level=3)
    _add_kv(doc, "Default reviewer", "Claude Sonnet 4.6")
    _add_kv(doc, "Security + Guided Apply", "Claude Opus 4.7")
    _add_kv(doc, "Deep Analysis (opt-in)", "OpenAI o3")
    _add_kv(doc, "Relevance gate + fast syntax", "Claude Haiku 4.5")
    _add_kv(doc, "Fallback provider", "GPT-5.5 (Anthropic outage) / Claude Sonnet 4.6 (OpenAI outage)")
    _add_kv(doc, "Cost cap", "Configurable per user and per day at gateway")


def main():
    src = "Genie_Platform_Documentation.docx"
    dst = "Genie_Platform_Documentation_Updated.docx"
    doc = Document(src)
    append_section(doc)
    doc.save(dst)
    print(f"Saved: {dst}")
    print(f"  (original {src} unchanged)")


if __name__ == "__main__":
    main()
