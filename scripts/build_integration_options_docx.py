"""Generate Integration_Options.docx — a client-facing comparison of Copilot API
vs. direct model API (Claude / OpenAI / OpenCode-style) for the Genie platform.

Run:  python scripts/build_integration_options_docx.py
Output: Integration_Options.docx at repo root.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BRAND_BLUE = RGBColor(0x1F, 0x3A, 0x68)
HEADER_GREY = RGBColor(0xE8, 0xEC, 0xF1)
BORDER_GREY = "BFBFBF"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color=BORDER_GREY, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def style_cell(cell, text, bold=False, size=10, color=None, shading=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if shading:
        set_cell_shading(cell, shading)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # Header row
    for i, h in enumerate(headers):
        style_cell(
            table.rows[0].cells[i],
            h,
            bold=True,
            size=10,
            color=RGBColor(0xFF, 0xFF, 0xFF),
            shading="1F3A68",
        )

    # Body rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            style_cell(
                table.rows[r].cells[c],
                val,
                bold=False,
                size=10,
                shading="F7F9FC" if r % 2 == 0 else None,
            )
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_BLUE
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def add_numbered(doc, text, size=11):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def build():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ===== Cover =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("INTEGRATION PATH OPTIONS")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND_BLUE

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Genie Platform")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = BRAND_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Copilot API vs. Direct Model Access — Limitations, Prerequisites, Effort")
    r.italic = True
    r.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Prepared by Bilvantis  •  Confidential  •  2025")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ===== About =====
    add_heading(doc, "About This Document", level=1)
    add_para(
        doc,
        "This document compares two production paths for the Genie platform beyond the current "
        "sandbox/PoC: continuing with the GitHub Copilot API, or migrating to direct model-provider "
        "APIs (Claude / OpenAI / OpenCode-style). Each option is evaluated across licensing, "
        "technical limitations, prerequisites, pros and cons, effort-to-production, and required "
        "tech stack.",
    )

    # ===== Section 1: Executive Summary =====
    add_heading(doc, "Section 1: Executive Summary", level=1)
    add_para(
        doc,
        "Two viable paths exist to move Genie beyond the current sandbox/PoC. Each has distinct "
        "legal, operational, and cost implications. This section compares them objectively so the "
        "client can make an informed decision on production direction.",
    )
    add_table(
        doc,
        headers=["Dimension", "Option A — GitHub Copilot API", "Option B — Direct Model APIs"],
        rows=[
            ["Status", "PoC-proven in current build", "Requires migration of LLM layer"],
            ["Licensing risk", "High — undocumented endpoint usage", "Low — fully licensed commercial APIs"],
            ["Effort to production", "4–6 months (legal-heavy)", "4–6 weeks (engineering-heavy)"],
            ["Recurring cost model", "Per-user Copilot seat", "Per-token consumption"],
            [
                "Model flexibility",
                "Restricted to Copilot catalogue",
                "Full model freedom (GPT-5.5, Claude Opus 4.7, Sonnet 4.6, o3, o4-mini, local)",
            ],
        ],
        col_widths=[4.0, 5.5, 6.0],
    )

    # ===== Section 2: Option A =====
    add_heading(doc, "Section 2: Option A — Continuing with GitHub Copilot API", level=1)

    add_heading(doc, "2.1  Current Implementation", level=2)
    add_para(
        doc,
        "The PoC uses GitHub's OAuth device-code flow to obtain a user access token, exchanges it "
        "for a short-lived session token, and calls https://api.githubcopilot.com/chat/completions "
        "with editor-identification headers (Editor-Version, Copilot-Integration-Id, "
        "Editor-Plugin-Version). The endpoint is OpenAI-compatible and supports tool calling, "
        "which is what the current LangGraph agent depends on.",
    )

    add_heading(doc, "2.2  Limitations (Production)", level=2)
    add_table(
        doc,
        headers=["#", "Limitation", "Impact"],
        rows=[
            [
                "1",
                "No official third-party API contract. The api.githubcopilot.com endpoint is intended for first-party GitHub editor integrations. The Copilot-Integration-Id header is not publicly documented for external partners.",
                "GitHub can revoke access, change headers, or rate-limit at any time without notice. No SLA.",
            ],
            [
                "2",
                "Terms of Service risk. GitHub's Acceptable Use policy prohibits reselling or redistributing Copilot capabilities through a third-party product without a partnership agreement.",
                "For enterprise deployment, legal sign-off is mandatory. A formal partnership with GitHub is the only risk-free path.",
            ],
            [
                "3",
                "Per-user Copilot licence required. Every developer using Genie must have an active GitHub Copilot subscription.",
                "Cost is paid twice: once to GitHub (per seat), once to the client organisation operating Genie.",
            ],
            [
                "4",
                "Undocumented rate limits. Per-user throttles are not published; heavy agentic use (multi-tool loops, whole-repo reviews) may trigger silent rate-limiting.",
                "Unpredictable latency spikes in production.",
            ],
            [
                "5",
                "Model catalogue lock-in. Only models GitHub routes through Copilot are available. No access to Claude Opus 4.7, o3, o4-mini, fine-tuned or local models.",
                "Cannot use best-in-class models where they matter (e.g. Opus for guided-apply).",
            ],
            [
                "6",
                "No enterprise audit trail. Copilot's telemetry is not exposed to the consuming application.",
                "Cannot produce per-call logs, token-accounting reports, or data-governance evidence required by regulated industries.",
            ],
            [
                "7",
                "Data residency limited. Requests route through GitHub-operated infrastructure; residency is controlled by GitHub's policies, not the client.",
                "Incompatible with clients requiring EU-only or APAC-only data handling guarantees.",
            ],
            [
                "8",
                "Session-token expiry. The session token is short-lived (minutes), requiring refresh logic and complicating long-running agent runs.",
                "Handled in PoC but increases failure surface.",
            ],
        ],
        col_widths=[0.8, 6.5, 6.2],
    )

    add_heading(doc, "2.3  Prerequisites", level=2)
    add_numbered(
        doc,
        "Legal — Written partnership or licensing agreement with GitHub confirming third-party API "
        "use is permitted for the intended scope. Without this, production use carries legal risk.",
    )
    add_numbered(doc, "Procurement — Volume licensing for GitHub Copilot Enterprise for all end-users.")
    add_numbered(
        doc,
        "Network — Allow-list api.githubcopilot.com and github.com through corporate proxy; "
        "CA-bundle handling for intercepting proxies (partially implemented).",
    )
    add_numbered(
        doc,
        "Compliance — Data Protection Impact Assessment if PII or regulated data could enter prompts.",
    )
    add_numbered(
        doc,
        "Monitoring — Client-side token-usage tracking and error-rate dashboards, since Copilot provides none.",
    )

    add_heading(doc, "2.4  Pros and Cons", level=2)
    add_para(doc, "Pros", bold=True)
    add_bullet(doc, "Zero migration cost — the PoC already works.")
    add_bullet(doc, "Developers authenticate once with their existing GitHub identity.")
    add_bullet(doc, "Access to multiple models through GitHub's routing without managing provider keys.")
    add_bullet(doc, "No separate billing infrastructure — Copilot licence covers LLM cost.")

    add_para(doc, "Cons", bold=True)
    add_bullet(doc, "Legal ambiguity is the single biggest blocker for enterprise sale.")
    add_bullet(doc, "Cannot pick the best model per task — critical for quality differentiation.")
    add_bullet(doc, "No SLA or support channel when things break.")
    add_bullet(doc, "Ties Genie's viability to GitHub's strategic decisions.")

    add_heading(doc, "2.5  Effort Estimate to Production", level=2)
    add_table(
        doc,
        headers=["Workstream", "Effort"],
        rows=[
            ["Legal partnership negotiation with GitHub", "3–6 months (blocking, out of engineering control)"],
            ["Hardening token management, retry, observability", "1 week"],
            ["Enterprise compliance documentation (DPIA, audit logs)", "2 weeks"],
            ["Rate-limit handling, graceful degradation", "1 week"],
            ["Total engineering", "~4 weeks"],
            ["Total realistic", "4–6 months including legal"],
        ],
        col_widths=[9.0, 6.5],
    )

    add_heading(doc, "2.6  Tech Stack (Option A)", level=2)
    add_bullet(doc, "Runtime: Node.js 20+ (VS Code extension host), Python 3.11+ (agent runtime)")
    add_bullet(doc, "Agent orchestration: LangGraph 0.2+, LangChain 0.3+")
    add_bullet(doc, "LLM client: langchain-openai pointed at Copilot proxy URL")
    add_bullet(doc, "Auth: GitHub OAuth device-code flow (already implemented)")
    add_bullet(doc, "Transport: HTTPS with corporate-proxy + PEM bundle support")
    add_bullet(doc, "Tooling: TypeScript 5.5, @vscode/vsce, Python httpx, certifi")
    add_bullet(doc, "No change from current PoC stack")

    # ===== Section 3: Option B =====
    add_heading(doc, "Section 3: Option B — Direct Model APIs (Claude / OpenAI / OpenCode-Style)", level=1)

    add_heading(doc, "3.1  Overview", level=2)
    add_para(
        doc,
        "This option replaces the Copilot proxy with direct, licensed API access to one or more LLM "
        "providers. 'OpenCode-style' here means the same architectural pattern used by open-source "
        "agent frameworks (OpenCode, Aider, Continue): a bring-your-own-key model where the extension "
        "ships without hard-coded LLM routing and the client configures their provider of choice.",
    )

    add_heading(doc, "3.2  Recommended Model Routing", level=2)
    add_table(
        doc,
        headers=["Task", "Recommended Model", "Justification"],
        rows=[
            [
                "Default review loop (Quality, Performance, OrgStd, CK Design)",
                "Claude Sonnet 4.6 or GPT-5.5",
                "Strongest tool-calling reliability; streams first tokens within 1–2s; ~1/5 cost of flagship tier.",
            ],
            [
                "Security review, Guided Apply (whole-file edits)",
                "Claude Opus 4.7",
                "Lowest hallucination rate on whole-file edits; strongest on subtle security reasoning.",
            ],
            [
                "Deep analysis mode (user-triggered)",
                "OpenAI o3 or o4-mini",
                "Best SWE-bench Verified scores at release (~70%); strong agentic tool use; trade higher first-token latency for rigour.",
            ],
            [
                "Relevance gate / quick syntax",
                "Claude Haiku 4.5 or GPT-5 mini",
                "Fast, cheap, sufficient for yes/no classification and trivial formatting.",
            ],
            [
                "Local / air-gapped deployments",
                "Llama 3.3 70B, Qwen 2.5 Coder 32B via Ollama / vLLM",
                "Only available in this option — not possible under Copilot.",
            ],
        ],
        col_widths=[5.0, 4.5, 6.0],
    )
    add_para(
        doc,
        "Note: o3 and o4-mini are not math-only models. OpenAI positioned them specifically as their "
        "most capable agentic-tool-using models at release (April 2025), with top-tier SWE-bench "
        "Verified scores. They are strong coding models with one tradeoff: higher first-token latency. "
        "In a developer-facing IDE, they should be gated behind an explicit 'Deep Analysis' action "
        "rather than the default flow. Cross-check benchmark numbers at "
        "https://openai.com/index/introducing-o3-and-o4-mini/ before quoting specific figures.",
        italic=True,
        size=10,
    )

    add_heading(doc, "3.3  Limitations", level=2)
    add_table(
        doc,
        headers=["#", "Limitation", "Impact"],
        rows=[
            [
                "1",
                "Client pays per token. Agentic loops re-send context on each turn.",
                "Cost proportional to usage; budget required. Estimated $0.05–$0.50 per full review depending on model mix.",
            ],
            [
                "2",
                "API key distribution. Each developer or the organisation centrally must hold valid keys for chosen providers.",
                "Requires secure key-distribution solution (Azure Key Vault, AWS Secrets Manager, or centralised gateway).",
            ],
            [
                "3",
                "Provider outages. Claude and OpenAI have had multi-hour outages in 2024–2025.",
                "Need fallback routing (e.g. GPT-5.5 if Claude is down) or graceful degradation.",
            ],
            [
                "4",
                "Per-provider quirks. Tool-call schemas, streaming chunk shapes, and rate limits differ between providers.",
                "Abstraction layer required; LangChain provides this but edge cases emerge.",
            ],
            [
                "5",
                "Data governance. Prompts sent to a third-party API.",
                "Requires DPA with each provider. Anthropic and OpenAI offer enterprise zero-retention contracts.",
            ],
            [
                "6",
                "First-token latency varies. Sonnet/GPT-5.5: 0.5–1.5s. o3/Opus: 3–10s.",
                "Default model must be chosen carefully for IDE feel.",
            ],
        ],
        col_widths=[0.8, 6.5, 6.2],
    )

    add_heading(doc, "3.4  Prerequisites", level=2)
    add_numbered(doc, "Commercial — Paid API accounts with chosen providers (Anthropic, OpenAI). Enterprise tier for zero-retention + DPA.")
    add_numbered(
        doc,
        "Infrastructure — Centralised LLM gateway (recommended): LiteLLM, Portkey, or custom Node.js proxy. "
        "Provides per-user quotas, logging, cost tracking, and provider fallback.",
    )
    add_numbered(doc, "Secrets — Key-management system for storing and rotating provider API keys.")
    add_numbered(
        doc,
        "Compliance — DPA signed with each provider. Data residency region selection "
        "(Anthropic: US/EU; OpenAI: US/EU/APAC via Azure OpenAI).",
    )
    add_numbered(doc, "Observability — Per-call logging of tokens, latency, model used, cost. Grafana or similar dashboard.")

    add_heading(doc, "3.5  Pros and Cons", level=2)
    add_para(doc, "Pros", bold=True)
    add_bullet(doc, "Fully licensed — no legal ambiguity; standard commercial contracts.")
    add_bullet(doc, "Best model per task — material quality improvement on security and guided-apply flows.")
    add_bullet(doc, "Predictable SLAs — Anthropic and OpenAI publish uptime targets and support channels.")
    add_bullet(doc, "Data residency and retention controls — enterprise offerings include zero-retention and region pinning.")
    add_bullet(doc, "Future-proof — can swap models, add local models, add new providers without re-architecture.")
    add_bullet(doc, "Observability — complete per-call audit trail owned by the client.")

    add_para(doc, "Cons", bold=True)
    add_bullet(doc, "Usage cost — clients pay per-token, not per-seat.")
    add_bullet(doc, "Migration effort — LLM layer must be abstracted behind a provider interface.")
    add_bullet(doc, "Key distribution complexity — solved by gateway but requires one-time setup.")
    add_bullet(doc, "Multiple vendor relationships — procurement, DPAs, and billing with 2–3 providers.")

    add_heading(doc, "3.6  Effort Estimate to Production", level=2)
    add_table(
        doc,
        headers=["Workstream", "Effort"],
        rows=[
            ["Replace Copilot proxy with provider-agnostic client abstraction", "3–5 days"],
            ["Centralised LLM gateway (LiteLLM or custom) with per-user keys", "1 week"],
            ["Per-provider DPA and zero-retention enrolment (legal + procurement)", "2–4 weeks (parallelisable)"],
            ["Cost tracking, logging, observability dashboard", "1 week"],
            ["Fallback routing + error handling", "3 days"],
            ["Model-selection UI and per-task model routing", "3 days"],
            ["Total engineering", "~3–4 weeks"],
            ["Total realistic", "4–6 weeks including procurement"],
        ],
        col_widths=[9.0, 6.5],
    )

    add_heading(doc, "3.7  Tech Stack (Option B)", level=2)
    add_bullet(doc, "Runtime: Node.js 20+ (extension), Python 3.11+ (agent runtime)")
    add_bullet(doc, "Agent orchestration: LangGraph 0.2+, LangChain 0.3+ (unchanged)")
    add_bullet(doc, "LLM clients: langchain-anthropic, langchain-openai, langchain-ollama (for local)")
    add_bullet(
        doc,
        "LLM gateway (recommended): LiteLLM or Portkey — unified OpenAI-compatible endpoint, "
        "routing, rate limiting, cost tracking",
    )
    add_bullet(doc, "Secrets management: Azure Key Vault / AWS Secrets Manager / HashiCorp Vault")
    add_bullet(doc, "Auth: API-key-per-provider via secure delivery to gateway, not shipped in extension")
    add_bullet(doc, "Observability: OpenTelemetry + Grafana (or equivalent) for latency, token, and cost dashboards")
    add_bullet(doc, "Transport: HTTPS with corporate-proxy + PEM bundle support (unchanged)")
    add_bullet(doc, "Optional — local models: Ollama or vLLM server for on-prem LLM hosting")

    # ===== Section 4: Head-to-Head =====
    add_heading(doc, "Section 4: Head-to-Head Summary", level=1)
    add_table(
        doc,
        headers=["Factor", "Option A (Copilot)", "Option B (Direct APIs)", "Winner"],
        rows=[
            ["Legal certainty", "Ambiguous", "Standard contracts", "B"],
            ["Model quality per task", "Locked catalogue", "Best-in-class per task", "B"],
            ["Recurring cost (light use, <50 reviews/dev/month)", "Full seat", "Pay-per-token", "B"],
            ["Recurring cost (heavy use, >500 reviews/dev/month)", "Capped at seat", "Scales linearly", "A"],
            ["SLA and support", "None", "Enterprise tier", "B"],
            ["Data governance and audit", "Opaque", "Full control", "B"],
            ["Engineering effort to production", "Minimal", "3–4 weeks", "A"],
            ["Time to signed commercial deal", "3–6 months", "4–6 weeks", "B"],
            ["Future flexibility (local models, new providers)", "None", "Unlimited", "B"],
        ],
        col_widths=[6.0, 4.0, 4.0, 1.5],
    )

    # ===== Section 5: Recommendation =====
    add_heading(doc, "Section 5: Recommendation", level=1)
    add_para(
        doc,
        "Proceed with Option B — Direct Model APIs — for production.",
        bold=True,
    )
    add_para(
        doc,
        "The PoC on Copilot has de-risked the agent architecture and proven the value proposition. "
        "For production rollout to enterprise clients, the legal ambiguity, lack of SLA, and "
        "model-catalogue restrictions of the Copilot path outweigh the short-term migration savings.",
    )
    add_para(doc, "Option B delivers:", bold=True)
    add_bullet(doc, "Commercial-grade licensing.")
    add_bullet(doc, "Best-in-class model quality where it matters (security, guided-apply).")
    add_bullet(doc, "Faster path to a signed enterprise contract — weeks, not months.")
    add_bullet(doc, "Strategic option to add local/on-prem models for regulated sectors.")

    add_para(doc, "Suggested phased transition:", bold=True)
    add_numbered(doc, "Week 1–2: abstract LLM client, add Anthropic + OpenAI drivers behind current LangGraph code.")
    add_numbered(doc, "Week 3: stand up LiteLLM gateway with centralised key management.")
    add_numbered(doc, "Week 4: observability, cost tracking, model-per-task routing configuration.")
    add_numbered(doc, "Parallel: procurement initiates DPAs with Anthropic and OpenAI.")

    doc.save("Integration_Options.docx")
    print("Saved: Integration_Options.docx")


if __name__ == "__main__":
    build()
